# -*- coding: utf-8 -*-
"""Generate Agent-friendly Knowledge Source files for the A4A hands-on (generic).

3 enriched Word docs (Security/Legacy/Platform) + 1 rich Excel (MarketCase).
Design goal: documents that a Copilot agent grounds on cleanly -> clear headings,
a TL;DR summary, simple flat tables, explicit FAQ Q/A pairs, defined terms.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTDIR = os.path.join(os.path.dirname(__file__), "KnowledgeSource")
os.makedirs(OUTDIR, exist_ok=True)

ACCENT = RGBColor(0xB1, 0x1F, 0x4B)
DARK = RGBColor(0x24, 0x24, 0x24)
MUTED = RGBColor(0x5C, 0x5C, 0x5C)
KFONT = "맑은 고딕"


def base_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = KFONT
    style.element.rPr.rFonts.set(qn("w:eastAsia"), KFONT)
    style.font.size = Pt(10.5)
    style.font.color.rgb = DARK
    return doc


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_run(run, size=None, bold=None, color=None):
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.name = KFONT
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), KFONT)


def title_block(doc, doc_id, title, subtitle, owner, updated, classification):
    p = doc.add_paragraph()
    _set_run(p.add_run(f"[{doc_id}]"), size=10, bold=True, color=ACCENT)
    h = doc.add_paragraph()
    _set_run(h.add_run(title), size=20, bold=True, color=DARK)
    s = doc.add_paragraph()
    _set_run(s.add_run(subtitle), size=11, color=MUTED)
    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["주관", "최종 수정", "보안등급", "용도"]
    vals = [owner, updated, classification, "Copilot Studio Knowledge Source"]
    for i in range(4):
        c0 = meta.rows[0].cells[i]
        c0.text = labels[i]
        shade_cell(c0, "F5F5F5")
        _set_run(c0.paragraphs[0].runs[0], size=8.5, bold=True)
        c1 = meta.rows[1].cells[i]
        c1.text = vals[i]
        _set_run(c1.paragraphs[0].runs[0], size=8.5)
    doc.add_paragraph()


def h2(doc, text):
    p = doc.add_paragraph()
    _set_run(p.add_run(text), size=13, bold=True, color=ACCENT)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def para(doc, text):
    p = doc.add_paragraph()
    _set_run(p.add_run(text), size=10.5)
    p.paragraph_format.space_after = Pt(2)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        _set_run(p.add_run(it), size=10)
        p.paragraph_format.space_after = Pt(1)


def summary_box(doc, items):
    """TL;DR box - agents read this first."""
    h2(doc, "한눈 요약 (TL;DR)")
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    shade_cell(cell, "FDF2F6")
    cell.text = ""
    first = True
    for it in items:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        _set_run(p.add_run("• " + it), size=9.5)
        first = False
    doc.add_paragraph()


def faq(doc, qa_pairs):
    h2(doc, "자주 묻는 질문 (FAQ)")
    for q, a in qa_pairs:
        pq = doc.add_paragraph()
        _set_run(pq.add_run("Q. " + q), size=10, bold=True, color=DARK)
        pq.paragraph_format.space_after = Pt(1)
        pa = doc.add_paragraph()
        _set_run(pa.add_run("A. " + a), size=10, color=MUTED)
        pa.paragraph_format.space_after = Pt(6)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = htext
        shade_cell(c, "B11F4B")
        _set_run(c.paragraphs[0].runs[0], size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            _set_run(cells[i].paragraphs[0].runs[0], size=9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def keywords(doc, kws):
    p = doc.add_paragraph()
    _set_run(p.add_run("키워드: "), size=8.5, bold=True, color=MUTED)
    _set_run(p.add_run(", ".join(kws)), size=8.5, color=MUTED)
    p.paragraph_format.space_after = Pt(2)


def usage_note(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    shade_cell(cell, "EEF4FB")
    cell.text = ""
    p = cell.paragraphs[0]
    _set_run(p.add_run("이 문서 사용법: "), size=9, bold=True, color=DARK)
    _set_run(p.add_run(text), size=9, color=DARK)
    doc.add_paragraph()


def footer_note(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    _set_run(p.add_run("※ 본 문서는 A4A 핸즈온 실습용 mock 데이터입니다. 실제 조직의 규정·시스템 정보가 아니며 교육 목적으로만 사용합니다."),
             size=8, color=MUTED)
    p.runs[0].italic = True


# ---------------- 1. Security_Rule ----------------
def build_security():
    doc = base_doc()
    title_block(doc, "Security_Rule",
                "생성형 AI · 에이전트 정보보안 가이드라인",
                "에이전트 설계 시 반드시 준수해야 할 데이터 보안 규정",
                "정보보안팀", "2026-05-30", "내부용")
    keywords(doc, ["정보등급", "데이터 분류", "개인정보", "DLP", "Shadow AI", "HITL", "보안 체크리스트", "승인 도구"])
    usage_note(doc, "에이전트가 다루려는 데이터의 보안 등급을 먼저 확인하고, 그 등급에서 허용되는 플랫폼·입력 방식만 설계에 반영하세요.")

    summary_box(doc, [
        "사내 데이터는 공개·내부·대외비·극비 4단계로 분류한다.",
        "대외비 이상 정보는 외부 공개 LLM(ChatGPT 등)에 입력 금지.",
        "승인된 도구는 Microsoft 365 Copilot, Copilot Studio(사내 테넌트), Azure OpenAI(사내 구독)뿐이다.",
        "외부 공유 산출물은 담당자 검토(HITL) 후 배포한다.",
    ])

    h2(doc, "1. 정보 등급 분류 체계")
    para(doc, "모든 사내 데이터는 아래 4단계로 분류되며, 등급에 따라 생성형 AI·에이전트 활용 범위가 결정됩니다.")
    table(doc,
          ["등급", "정의", "예시", "생성형 AI 입력", "권장 플랫폼"],
          [
              ["공개(Public)", "외부 공개 가능 정보", "보도자료, 홈페이지 게시물", "허용", "제한 없음"],
              ["내부(Internal)", "사내 한정 정보", "내부 공지, 일반 회의록", "사내 승인 AI만 허용", "M365 Copilot / Agent Builder"],
              ["대외비(Confidential)", "유출 시 경영 영향", "미공개 실적, 가격 정책, R&D 진행", "외부 AI 입력 금지", "Copilot Studio(사내 테넌트)"],
              ["극비(Restricted)", "유출 시 중대 손실", "특허 전략, M&A, 핵심 공정 레시피", "전면 금지", "에이전트 활용 불가"],
          ],
          widths=[1.1, 1.6, 1.9, 1.3, 1.5])

    h2(doc, "2. 개인정보 보호 규칙")
    bullets(doc, [
        "고객·구성원의 개인정보(이름, 연락처, 사번, 주민번호 등)는 외부 생성형 AI에 입력 금지.",
        "에이전트가 개인정보를 다룰 경우 Microsoft 365 테넌트 내(DLP 적용 환경)에서만 처리.",
        "테스트·데모 데이터는 반드시 가명/익명 처리된 mock 데이터 사용.",
        "개인정보가 포함된 Knowledge Source는 접근 권한을 부서 단위로 제한.",
    ])

    h2(doc, "3. 외부 전송 · 외부 도구 사용 규칙")
    bullets(doc, [
        "대외비 이상 정보는 사외 클라우드·공개 LLM(ChatGPT 등 외부 서비스)에 입력 절대 금지.",
        "승인된 도구: Microsoft 365 Copilot, Copilot Studio(사내 테넌트), Azure OpenAI(사내 구독).",
        "미승인 외부 AI 도구 사용(Shadow AI) 적발 시 정보보안 위반으로 처리.",
        "에이전트의 외부 웹 검색 기능은 내부 자료가 질의문에 포함되지 않도록 설계.",
    ])

    h2(doc, "4. 에이전트 설계 시 보안 체크리스트")
    table(doc,
          ["항목", "확인 내용", "필수 여부"],
          [
              ["데이터 등급 확인", "에이전트가 다루는 데이터의 최고 등급 식별", "필수"],
              ["입력 검증", "사용자 입력에 극비 정보 포함 시 차단/마스킹", "필수"],
              ["권한 범위", "Knowledge Source 접근 권한 최소화(필요 부서만)", "필수"],
              ["로그·감사", "대화/처리 이력 로그 보관 및 주기적 점검", "권장"],
              ["출력 검토(HITL)", "외부 공유 산출물은 담당자 승인 후 배포", "필수"],
          ],
          widths=[1.6, 3.6, 1.2])

    faq(doc, [
        ("대외비 자료를 ChatGPT에 넣어 요약해도 되나요?",
         "안 됩니다. 대외비 이상 정보는 외부 공개 LLM 입력이 전면 금지이며, Copilot Studio(사내 테넌트) 같은 승인 도구만 사용해야 합니다."),
        ("에이전트가 개인정보를 다뤄야 하면 어떻게 설계하나요?",
         "Microsoft 365 테넌트 내 DLP 적용 환경에서만 처리하고, Knowledge Source 접근 권한을 필요한 부서로 최소화하세요."),
        ("내부 등급 자료로 만든 에이전트는 어떤 플랫폼이 적합한가요?",
         "내부 등급은 사내 승인 AI(M365 Copilot, Agent Builder)에서 활용 가능합니다. 대외비가 섞이면 Copilot Studio로 올려야 합니다."),
        ("외부 공유용 결과물은 바로 보내도 되나요?",
         "아니요. 외부로 나가는 산출물은 담당자 검토(HITL) 승인 후 배포해야 합니다."),
    ])

    h2(doc, "5. 위반 시 조치")
    para(doc, "보안 규정 위반 사항 발견 시 즉시 정보보안팀(security@example.com)에 통보하며, 위반 수준에 따라 에이전트 사용 중단·재설계·징계 절차가 적용됩니다.")

    footer_note(doc)
    out = os.path.join(OUTDIR, "Security_Rule.docx")
    doc.save(out)
    return out


# ---------------- 2. Legacy_System_Map ----------------
def build_legacy():
    doc = base_doc()
    title_block(doc, "Legacy_System_Map",
                "사내 시스템 · 데이터 연계 맵",
                "에이전트 연동 가능 시스템과 데이터 접근성 정보",
                "시스템개발팀", "2026-05-28", "내부용")
    keywords(doc, ["사내 시스템", "ERP", "그룹웨어", "SharePoint", "API 연동", "데이터 접근성", "커넥터", "그라운딩"])
    usage_note(doc, "에이전트가 어떤 사내 데이터를 쓸 수 있는지 판단할 때 참조하세요. 접근성 '높음' 시스템부터 연동을 권장합니다.")

    summary_box(doc, [
        "SharePoint Online은 접근성이 가장 높아 문서 기반 에이전트의 1순위 데이터 소스다.",
        "SAP·LIMS·PLM은 제한적 연동(중)으로 IT 승인과 커넥터 구성이 필요하다.",
        "그룹웨어·MES는 표준 API가 없어(낮음) 추출 후 SharePoint 업로드 방식을 권장한다.",
        "운영 시스템에 대한 직접 쓰기(write)는 금지하고 읽기 전용으로 설계한다.",
    ])

    h2(doc, "1. 핵심 사내 시스템 개요")
    para(doc, "일반 LLM이 알 수 없는 사내 시스템 정보입니다. 에이전트 설계 시 연동 가능성과 접근 방식을 아래 표로 판단합니다.")
    table(doc,
          ["시스템", "용도", "보유 데이터", "연동 방식", "갱신 주기", "접근성"],
          [
              ["SAP S/4HANA", "전사 ERP(재무·구매·생산)", "전표, 구매오더, 생산실적", "OData API(읽기 위주)", "실시간", "중 (승인 필요)"],
              ["그룹웨어", "전자결재·메일·게시판", "결재 문서, 공지", "표준 API 미지원(RPA 보완)", "수시", "낮음"],
              ["SharePoint Online", "문서·지식 저장소", "보고서, 양식, 매뉴얼", "Microsoft Graph API", "수시", "높음"],
              ["MES", "생산실행시스템", "공정 데이터, 설비 가동", "사내 IF 서버 경유", "실시간", "낮음 (현장 한정)"],
              ["LIMS", "품질·시험 관리", "시험 성적서, 규격", "DB 직접조회(제한)", "일 배치", "중"],
              ["PLM", "제품수명주기관리", "BOM, 스펙, 설계변경", "전용 커넥터", "수시", "중"],
          ],
          widths=[1.1, 1.5, 1.5, 1.5, 0.9, 1.1])

    h2(doc, "2. 데이터 접근성 등급 정의")
    bullets(doc, [
        "높음: Graph/REST API 표준 지원, 에이전트 직접 연동 권장.",
        "중: 제한적 API·DB 조회 가능, IT 승인 및 커넥터 구성 필요.",
        "낮음: 표준 연동 불가, 수동 업로드 또는 RPA로 단기 Knowledge Source화 권장.",
    ])

    h2(doc, "3. 권장 연동 패턴")
    table(doc,
          ["시나리오", "권장 패턴"],
          [
              ["문서 기반 질의응답", "SharePoint 문서 → Copilot Studio Knowledge Source 등록"],
              ["실적·수치 조회", "SAP OData → Power Automate 커넥터 → Dataverse 캐싱"],
              ["품질 성적 조회", "LIMS 추출 데이터 → 주기적 SharePoint 업로드 → 그라운딩"],
              ["결재·메일 연계", "그룹웨어는 직접 연동 대신 산출물 SharePoint 저장 후 활용"],
          ],
          widths=[2.4, 4.0])

    h2(doc, "4. 시스템별 에이전트 활용 예시")
    table(doc,
          ["시스템", "가능한 에이전트 업무", "연동 난이도"],
          [
              ["SharePoint Online", "사내 규정·매뉴얼 Q&A, 보고서 초안 작성", "쉬움"],
              ["SAP S/4HANA", "구매오더 현황 조회, 생산실적 요약", "보통"],
              ["LIMS", "시험 성적서 검색, 규격 적합 여부 확인", "보통"],
              ["MES", "설비 가동률 일일 리포트(추출 데이터 기반)", "어려움"],
          ],
          widths=[1.6, 3.6, 1.2])

    h2(doc, "5. 연동 시 주의사항")
    bullets(doc, [
        "운영 시스템(SAP·MES) 직접 쓰기(write)는 원칙적으로 금지, 읽기 전용으로 설계.",
        "대용량 조회는 야간 배치/캐싱으로 분산하여 운영 시스템 부하 방지.",
        "시스템 연동은 IT 부서 사전 승인 및 테스트 환경 검증 후 운영 반영.",
    ])

    faq(doc, [
        ("문서 기반 에이전트를 가장 빠르게 만들려면 어떤 시스템을 쓰나요?",
         "SharePoint Online이 접근성 '높음'이라 가장 빠릅니다. 문서를 그대로 Copilot Studio Knowledge Source로 등록하면 됩니다."),
        ("SAP 데이터를 실시간으로 에이전트에 붙일 수 있나요?",
         "직접 실시간 연동보다 OData를 Power Automate 커넥터로 가져와 Dataverse에 캐싱하는 패턴을 권장합니다. 운영 부하를 줄일 수 있습니다."),
        ("그룹웨어 결재 문서를 에이전트가 읽게 하려면?",
         "그룹웨어는 표준 API가 없어 직접 연동이 어렵습니다. 산출물을 SharePoint에 저장한 뒤 그 문서를 그라운딩하세요."),
    ])

    footer_note(doc)
    out = os.path.join(OUTDIR, "Legacy_System_Map.docx")
    doc.save(out)
    return out


# ---------------- 3. Platform_Rule ----------------
def build_platform():
    doc = base_doc()
    title_block(doc, "Platform_Rule",
                "에이전트 제작 플랫폼 선정 룰셋",
                "직접 제작 vs A4A 활용 · 플랫폼별 적용 판단 기준",
                "People & AI팀", "2026-05-29", "내부용")
    keywords(doc, ["Agent Builder", "Copilot Studio", "Azure AI", "플랫폼 선정", "노코드", "HITL", "데이터 민감도", "의사결정"])
    usage_note(doc, "에이전트 아이디어를 받으면 데이터 민감도·통합 복잡도·사용자 규모·HITL 필요성 4축으로 플랫폼을 추천하세요.")

    summary_box(doc, [
        "개인 생산성·빠른 프로토타입은 Agent Builder(노코드, 5~10분).",
        "부서 단위 업무 자동화·사내 시스템 연동·승인 흐름은 Copilot Studio.",
        "대규모·복잡 커스텀 통합만 Azure AI 자체 개발을 검토한다.",
        "대외비 이상 데이터나 HITL이 필요하면 Agent Builder가 아닌 Copilot Studio 이상을 선택한다.",
    ])

    h2(doc, "1. 플랫폼 선택지 비교")
    table(doc,
          ["플랫폼", "특징", "적합 상황", "제작 난이도"],
          [
              ["Agent Builder (M365)", "노코드, 5~10분, 단일 화면", "개인 생산성, 빠른 프로토타입", "낮음"],
              ["Copilot Studio", "토픽·Flow·Knowledge·HITL", "부서 단위 업무 자동화", "중"],
              ["Azure AI / 자체 개발", "완전 커스텀, 코드 기반", "대규모·복잡 통합", "높음"],
          ],
          widths=[1.7, 2.0, 1.8, 1.0])

    h2(doc, "2. 판단 기준 (4대 축)")
    table(doc,
          ["기준", "낮음 → Agent Builder", "높음 → Copilot Studio 이상"],
          [
              ["데이터 민감도", "공개·내부 등급", "대외비 이상, DLP 필요"],
              ["통합 복잡도", "웹 검색·문서 생성 수준", "사내 시스템 연동 다수"],
              ["사용자 규모", "개인·소수", "부서·전사 배포"],
              ["HITL 필요성", "승인 절차 불필요", "다단계 승인·감사 필요"],
          ],
          widths=[1.6, 2.4, 2.4])

    h2(doc, "3. 의사결정 플로우")
    bullets(doc, [
        "Q1. 사내 시스템(SAP·LIMS 등) 연동이 필요한가? → 예: Copilot Studio 이상.",
        "Q2. 대외비 이상 데이터를 다루는가? → 예: Copilot Studio(사내 테넌트, DLP) 권장.",
        "Q3. 승인·감사 등 HITL 흐름이 필요한가? → 예: Copilot Studio(Flow + 승인).",
        "Q4. 위 모두 아니오 → Agent Builder로 빠르게 제작.",
        "Q5. 표준 플랫폼으로 불가능한 대규모 통합 → Azure AI 자체 개발 검토.",
    ])

    h2(doc, "4. 업무 유형별 플랫폼 추천")
    table(doc,
          ["업무 유형", "추천 플랫폼", "이유"],
          [
              ["개인 메일·문서 요약", "Agent Builder", "공개·내부 데이터, 단일 사용자"],
              ["부서 FAQ 봇(사내 규정)", "Copilot Studio", "사내 지식 그라운딩 + 다수 사용자"],
              ["계약 위험조항 검출", "Copilot Studio", "대외비 데이터 + 담당자 승인(HITL)"],
              ["전사 통합 데이터 분석", "Azure AI", "대규모·복잡 시스템 통합"],
          ],
          widths=[2.0, 1.8, 2.6])

    h2(doc, "5. A4A 활용 권장 시나리오")
    para(doc, "다음의 경우 직접 제작보다 A4A(에이전트 제작 코치)를 통해 설계서를 먼저 받는 것을 권장합니다.")
    bullets(doc, [
        "처음 에이전트를 만들어보는 구성원(기획·프롬프트 작성 경험 부족).",
        "어떤 플랫폼이 적합한지 판단이 서지 않는 경우.",
        "사내 보안·시스템 제약을 설계에 반영해야 하는 경우.",
    ])

    faq(doc, [
        ("노코드로 5분 만에 만들고 싶은데 어떤 플랫폼을 쓰나요?",
         "Agent Builder(M365)가 적합합니다. 단, 공개·내부 등급 데이터에 개인·소수 사용자 범위일 때만 권장합니다."),
        ("부서 전체가 쓰는 사내 규정 FAQ 봇은요?",
         "Copilot Studio를 쓰세요. 사내 지식 그라운딩(Knowledge)과 다수 사용자 배포, 필요 시 승인 흐름까지 지원합니다."),
        ("대외비 데이터를 다루는데 Agent Builder로 만들어도 되나요?",
         "권장하지 않습니다. 대외비 이상이거나 HITL이 필요하면 Copilot Studio 이상(사내 테넌트, DLP)을 선택하세요."),
    ])

    footer_note(doc)
    out = os.path.join(OUTDIR, "Platform_Rule.docx")
    doc.save(out)
    return out


# ---------------- 4. MarketCase_Knowledge (Excel) ----------------
def build_marketcase():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()

    accent = "B11F4B"
    head_fill = PatternFill("solid", fgColor=accent)
    head_font = Font(name=KFONT, bold=True, color="FFFFFF", size=10)
    body_font = Font(name=KFONT, size=10)
    good_fill = PatternFill("solid", fgColor="EAF6EE")
    bad_fill = PatternFill("solid", fgColor="FCEBEC")
    thin = Side(style="thin", color="DDDDDD")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # ---- Sheet 1: 사례목록 ----
    ws = wb.active
    ws.title = "사례목록"
    headers = ["사례ID", "구분", "본부", "업무유형", "에이전트명", "사용플랫폼",
               "데이터등급", "HITL적용", "반복빈도", "핵심성과지표(KPI)", "정량효과",
               "성공/실패요인", "핵심교훈"]
    rows = [
        ["C001", "우수", "마케팅본부", "시장조사", "MI 자동화 에이전트", "Copilot Studio", "내부", "Y", "주간", "리서치 소요시간", "4시간→30분", "명확한 단일 업무 정의", "범위를 좁히면 효과가 크다"],
        ["C002", "우수", "경영지원본부", "계약검토", "계약 검토 어시스턴트", "Copilot Studio", "대외비", "Y", "수시", "법무 검토 리드타임", "40% 단축", "사내 표준계약 그라운딩", "대외비는 사내 테넌트로"],
        ["C003", "우수", "생산본부", "보고서작성", "주간보고 생성기", "Copilot Studio", "내부", "Y", "주간", "보고서 작성시간", "3시간 절감", "실적 데이터 자동 연결", "데이터 연동이 핵심"],
        ["C004", "우수", "마케팅본부", "콘텐츠생성", "캠페인 카피 코치", "Agent Builder", "공개", "N", "일간", "초안 생산량", "2배 증가", "가벼운 범위로 빠른 제작", "노코드로 충분한 영역"],
        ["C005", "우수", "경영지원본부", "FAQ응대", "사내 규정 FAQ 봇", "Copilot Studio", "내부", "N", "상시", "문의 응대율", "1차 응대 70%", "사내 규정 문서 그라운딩", "지식 소스 품질이 답을 만든다"],
        ["C006", "우수", "연구개발본부", "문헌요약", "논문·특허 요약 에이전트", "Copilot Studio", "내부", "Y", "주간", "문헌 검토시간", "주 5시간 절감", "출처 명시 요약 설계", "근거 링크가 신뢰를 높인다"],
        ["C007", "우수", "품질본부", "성적서조회", "시험성적 조회 봇", "Copilot Studio", "내부", "N", "일간", "성적서 검색시간", "건당 10분→1분", "LIMS 추출 데이터 활용", "추출·업로드 패턴이 현실적"],
        ["C008", "우수", "인사본부", "온보딩안내", "신입 온보딩 가이드", "Agent Builder", "내부", "N", "상시", "HR 반복문의", "30% 감소", "FAQ형 단일 목적", "반복 질문부터 자동화"],
        ["C009", "우수", "재무본부", "비용분석", "월마감 점검 에이전트", "Copilot Studio", "대외비", "Y", "월간", "마감 점검시간", "반나절 절감", "체크리스트 기반 검증", "사람 검토(HITL)와 병행"],
        ["C010", "우수", "영업본부", "제안서작성", "RFP 제안 초안 봇", "Copilot Studio", "내부", "Y", "수시", "제안서 초안시간", "50% 단축", "과거 우수제안 그라운딩", "좋은 예시가 좋은 출력을 만든다"],
        ["C011", "우수", "구매본부", "현황조회", "구매오더 현황 봇", "Copilot Studio", "내부", "N", "일간", "현황 취합시간", "1시간→5분", "SAP OData 캐싱 연동", "운영 시스템은 읽기 전용"],
        ["C012", "우수", "마케팅본부", "경쟁분석", "경쟁사 동향 브리핑", "Copilot Studio", "내부", "Y", "주간", "브리핑 준비시간", "주 3시간 절감", "신뢰 가능한 출처 한정", "외부 검색 범위를 통제"],
        ["C013", "실패", "전사", "만능봇", "무엇이든 봇", "Copilot Studio", "내부", "N", "-", "실사용률", "제작 후 미사용", "범위 과대·업무와 무관", "작은 반복 업무부터 시작"],
        ["C014", "실패", "영업본부", "보고서작성", "실적 자동보고", "Agent Builder", "대외비", "N", "-", "정확도", "수치 환각 발생", "그라운딩 없는 자유생성", "Knowledge Source 필수 연결"],
        ["C015", "실패", "연구개발본부", "외부질의", "리서치 도우미", "외부 공개 LLM", "대외비", "N", "-", "보안 사고", "대외비 유출 위험", "승인 안 된 외부 AI(Shadow AI)", "승인 도구·등급 사전 확인"],
        ["C016", "실패", "경영지원본부", "계약검토", "초안 자동승인", "Copilot Studio", "대외비", "N", "-", "검토 누락", "오류 조항 통과", "사람 검토(HITL) 생략", "중요 결정엔 HITL 필수"],
        ["C017", "실패", "생산본부", "설비분석", "실시간 설비봇", "Azure AI", "내부", "N", "-", "운영 안정성", "운영 시스템 부하", "MES 직접 대량 조회", "야간 배치·캐싱으로 분산"],
        ["C018", "실패", "인사본부", "개인정보처리", "직원 응대봇", "외부 공개 LLM", "대외비", "N", "-", "컴플라이언스", "개인정보 입력 위반", "DLP 없는 외부 환경 사용", "개인정보는 사내 테넌트만"],
        ["C019", "실패", "마케팅본부", "콘텐츠생성", "트렌드 자동포스팅", "Agent Builder", "공개", "N", "-", "품질", "사실 오류 게시", "출력 검토 없이 자동발행", "외부 공개물은 검토 후 배포"],
        ["C020", "실패", "재무본부", "예측분석", "예산 예측봇", "Copilot Studio", "대외비", "N", "-", "신뢰도", "근거 없는 추정", "지식 소스 미연결", "정량 근거를 반드시 연결"],
        ["C021", "우수", "물류본부", "재고조회", "재고 현황 어시스턴트", "Copilot Studio", "내부", "N", "일간", "재고 확인시간", "건당 8분 절감", "단일 목적·명확한 질의", "좁은 업무가 빠르게 정착"],
        ["C022", "우수", "고객지원본부", "응대지원", "상담 가이드 봇", "Copilot Studio", "내부", "Y", "상시", "평균 응대시간", "20% 단축", "상담 매뉴얼 그라운딩", "상담원 보조로 시작"],
        ["C023", "실패", "전사", "교육봇", "AI 교육 도우미", "Agent Builder", "내부", "N", "-", "지속 사용률", "초기 사용 후 방치", "성과 지표 미정의", "측정 지표를 처음부터 설계"],
        ["C024", "우수", "법무본부", "규정검색", "사내 규정 검색봇", "Copilot Studio", "대외비", "Y", "상시", "규정 검색시간", "건당 15분→2분", "권한 분리·부서 한정 접근", "민감 지식은 권한 최소화"],
    ]
    ws.append(headers)
    for j, _ in enumerate(headers, 1):
        c = ws.cell(row=1, column=j)
        c.fill = head_fill
        c.font = head_font
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = border
    for i, row in enumerate(rows, start=2):
        ws.append(row)
        is_good = row[1] == "우수"
        for j in range(1, len(headers) + 1):
            c = ws.cell(row=i, column=j)
            c.font = body_font
            c.border = border
            c.alignment = Alignment(vertical="center", wrap_text=True)
            c.fill = good_fill if is_good else bad_fill
    widths = [7, 6, 12, 10, 20, 15, 9, 8, 8, 16, 14, 22, 24]
    from openpyxl.utils import get_column_letter
    for idx, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = w
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    ws.row_dimensions[1].height = 28

    # ---- Sheet 2: 필드정의 (Data Dictionary) ----
    ws2 = wb.create_sheet("필드정의")
    dd_head = ["컬럼명", "설명", "값 예시"]
    dd_rows = [
        ["사례ID", "사례 고유 식별자", "C001"],
        ["구분", "우수 사례 / 실패 사례", "우수, 실패"],
        ["본부", "사례가 발생한 조직 단위", "마케팅본부, 생산본부 등"],
        ["업무유형", "자동화한 업무의 종류", "시장조사, 계약검토, 보고서작성"],
        ["에이전트명", "제작된 에이전트 이름", "MI 자동화 에이전트"],
        ["사용플랫폼", "제작에 사용한 플랫폼", "Agent Builder, Copilot Studio, Azure AI"],
        ["데이터등급", "다룬 데이터의 최고 보안 등급", "공개, 내부, 대외비"],
        ["HITL적용", "담당자 검토·승인 단계 포함 여부", "Y, N"],
        ["반복빈도", "업무가 반복되는 주기", "상시, 일간, 주간, 월간"],
        ["핵심성과지표(KPI)", "효과를 측정한 지표", "리서치 소요시간 등"],
        ["정량효과", "지표의 개선 수치", "4시간→30분, 40% 단축"],
        ["성공/실패요인", "결과를 만든 핵심 원인", "단일 업무 정의, 그라운딩 누락 등"],
        ["핵심교훈", "설계 시 적용할 교훈(Lessons Learned)", "범위를 좁혀라 등"],
    ]
    ws2.append(dd_head)
    for j in range(1, len(dd_head) + 1):
        c = ws2.cell(row=1, column=j)
        c.fill = head_fill
        c.font = head_font
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = border
    for i, row in enumerate(dd_rows, start=2):
        ws2.append(row)
        for j in range(1, len(dd_head) + 1):
            c = ws2.cell(row=i, column=j)
            c.font = body_font
            c.border = border
            c.alignment = Alignment(vertical="center", wrap_text=True)
    for idx, w in enumerate([18, 40, 30], 1):
        ws2.column_dimensions[get_column_letter(idx)].width = w
    ws2.freeze_panes = "A2"

    # ---- Sheet 3: 설계교훈 요약 ----
    ws3 = wb.create_sheet("설계교훈")
    le_head = ["주제", "교훈", "근거 사례"]
    le_rows = [
        ["범위", "'모든 것을 하는 에이전트'보다 한 가지 업무를 확실히 정의한다", "C001, C013"],
        ["데이터", "사내 Knowledge Source 없이는 품질·신뢰를 확보할 수 없다", "C003, C014, C020"],
        ["사람", "외부 공유·중요 결정은 반드시 담당자 검토(HITL)를 거친다", "C002, C016, C019"],
        ["측정", "절감 시간·처리 건수 등 성과 지표를 설계 단계부터 정의한다", "C008, C023"],
        ["보안", "데이터 등급과 플랫폼 적합성을 가장 먼저 확인한다", "C015, C018, C024"],
        ["운영", "운영 시스템은 읽기 전용·배치/캐싱으로 부하를 막는다", "C011, C017"],
    ]
    ws3.append(le_head)
    for j in range(1, len(le_head) + 1):
        c = ws3.cell(row=1, column=j)
        c.fill = head_fill
        c.font = head_font
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = border
    for i, row in enumerate(le_rows, start=2):
        ws3.append(row)
        for j in range(1, len(le_head) + 1):
            c = ws3.cell(row=i, column=j)
            c.font = body_font
            c.border = border
            c.alignment = Alignment(vertical="center", wrap_text=True)
    for idx, w in enumerate([12, 50, 16], 1):
        ws3.column_dimensions[get_column_letter(idx)].width = w
    ws3.freeze_panes = "A2"

    # ---- Sheet 4: 안내 ----
    ws4 = wb.create_sheet("안내")
    notes = [
        "MarketCase_Knowledge — 사내 AI Market 우수·실패 사례 데이터 (A4A 핸즈온 실습용)",
        "",
        "이 파일 사용법: A4A가 에이전트 설계를 추천할 때 참조하는 사례 데이터입니다.",
        "'사례목록' 시트에서 본부·업무유형·플랫폼·데이터등급으로 필터링하여",
        "유사한 우수 사례는 참고하고, 실패 사례의 교훈은 회피하도록 설계에 반영하세요.",
        "",
        "시트 구성:",
        "  · 사례목록 — 24건의 우수/실패 사례 (필터·정렬 가능)",
        "  · 필드정의 — 각 컬럼의 의미와 값 예시(Data Dictionary)",
        "  · 설계교훈 — 사례에서 도출한 핵심 교훈(Lessons Learned)",
        "",
        "※ 본 데이터는 A4A 핸즈온 실습용 mock 데이터입니다.",
        "   실제 조직의 사례가 아니며 교육 목적으로만 사용합니다.",
    ]
    for i, line in enumerate(notes, start=1):
        c = ws4.cell(row=i, column=1, value=line)
        c.font = Font(name=KFONT, size=11, bold=(i == 1))
    ws4.column_dimensions["A"].width = 90

    # order: 안내 first is friendlier, but keep 사례목록 first for data-first agents
    out = os.path.join(OUTDIR, "MarketCase_Knowledge.xlsx")
    wb.save(out)
    return out


if __name__ == "__main__":
    outs = [build_security(), build_legacy(), build_platform(), build_marketcase()]
    for o in outs:
        print("saved:", o, os.path.getsize(o), "bytes")
    # remove old docx version of MarketCase if present
    old = os.path.join(OUTDIR, "MarketCase_Knowledge.docx")
    if os.path.exists(old):
        os.remove(old)
        print("removed old:", old)

# -*- coding: utf-8 -*-
"""A4A 설계서 문서 레이아웃(.docx) 생성기.
Copilot Studio 프롬프트 도구의 'Document(문서) 출력' 기능에 업로드하는 레이아웃.
필드는 {{필드명}} 형식(공백 없음). AI가 각 필드를 채워 Word 설계서를 생성한다.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTDIR = os.path.join(os.path.dirname(__file__), "KnowledgeSource")
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x55, 0x5F, 0x6B)


def set_cell_bg(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def style_base(doc):
    st = doc.styles['Normal']
    st.font.name = '맑은 고딕'
    st.font.size = Pt(10.5)
    from docx.oxml.ns import qn
    st.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    return p


def kv_table(doc, rows):
    """rows: list of (label, field_placeholder)"""
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    for label, val in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(4.2)
        cells[1].width = Cm(11.5)
        set_cell_bg(cells[0], 'EEF2F7')
        lp = cells[0].paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.color.rgb = GRAY
        cells[1].paragraphs[0].add_run(val)
    return table


def build():
    doc = Document()
    style_base(doc)
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run('AI 에이전트 제작 설계서')
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('A4A · Agent for Agent 인터뷰 결과 기반 설계서')
    sr.font.size = Pt(10)
    sr.font.color.rgb = GRAY

    doc.add_paragraph()

    # 1. 개요
    add_heading(doc, '1. 에이전트 개요')
    kv_table(doc, [
        ('에이전트 이름', '{{에이전트이름}}'),
        ('한 줄 목적', '{{한줄목적}}'),
        ('추천 플랫폼', '{{추천플랫폼}}'),
        ('선정 이유', '{{선정이유}}'),
    ])

    # 2. 지침 초안
    add_heading(doc, '2. 지침(Instructions) 초안')
    doc.add_paragraph('{{지침초안}}')

    # 3. 참조 자료
    add_heading(doc, '3. 필요한 참조 자료(Knowledge Source)')
    doc.add_paragraph('{{참조자료목록}}')

    # 4. 테스트 시나리오
    add_heading(doc, '4. 테스트 시나리오')
    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    hdr = t.add_row().cells
    set_cell_bg(hdr[0], '1F3A5F')
    set_cell_bg(hdr[1], '1F3A5F')
    for c, txt in zip(hdr, ['구분', '시나리오']):
        r = c.paragraphs[0].add_run(txt)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hdr[0].width = Cm(3.0)
    hdr[1].width = Cm(12.7)
    for label, field in [('시나리오 1', '{{테스트시나리오1}}'),
                         ('시나리오 2', '{{테스트시나리오2}}'),
                         ('시나리오 3', '{{테스트시나리오3}}')]:
        cells = t.add_row().cells
        set_cell_bg(cells[0], 'EEF2F7')
        lr = cells[0].paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.color.rgb = GRAY
        cells[1].paragraphs[0].add_run(field)

    # 5. 보안 체크리스트
    add_heading(doc, '5. 보안 체크리스트 확인 결과')
    doc.add_paragraph('{{보안체크리스트}}')

    # 푸터 안내
    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run('※ 이 문서는 Copilot Studio 프롬프트 도구의 문서 출력용 레이아웃입니다. '
                      '{{ }} 안의 필드는 에이전트가 인터뷰 내용을 바탕으로 자동으로 채웁니다.')
    nr.font.size = Pt(8.5)
    nr.font.color.rgb = GRAY

    os.makedirs(OUTDIR, exist_ok=True)
    out = os.path.join(OUTDIR, 'A4A_Design_Layout.docx')
    doc.save(out)
    print('saved:', out)


if __name__ == '__main__':
    build()
